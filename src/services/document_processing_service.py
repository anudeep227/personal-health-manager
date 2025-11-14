"""
Document Processing Service for Health Management App
Handles various document types: PDF, DOCX, images, medical reports, ECG data
"""

import os
import io
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import logging

# Core document processing
try:
    import PyPDF2
    from docx import Document
    import pytesseract
    import cv2
    import numpy as np
    from PIL import Image, ImageEnhance, ImageFilter
    PDF_AVAILABLE = True
except ImportError as e:
    PDF_AVAILABLE = False
    # Create dummy numpy for type hints when not available
    class np:
        ndarray = object
    logging.warning(f"Some document processing libraries not available: {e}")

# Advanced processing (optional)
try:
    import pdfplumber
    import easyocr
    import mammoth
    ADVANCED_PROCESSING = True
except ImportError:
    ADVANCED_PROCESSING = False

from src.utils.config import Config


class DocumentProcessingService:
    """Service for processing various document types and extracting text/data"""
    
    def __init__(self):
        self.config = Config()
        self.supported_formats = {
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'text': ['.txt', '.rtf'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'],
            'medical': ['.dcm', '.dicom']  # Medical imaging formats
        }
        
        # Initialize OCR readers
        self.ocr_reader = None
        self.setup_ocr()
        
        # Medical document patterns
        self.medical_patterns = {
            'ecg': ['electrocardiogram', 'ecg', 'ekg', 'heart rhythm', 'cardiac'],
            'blood_test': ['blood', 'hemoglobin', 'glucose', 'cholesterol', 'CBC'],
            'prescription': ['rx', 'prescription', 'medication', 'dosage', 'mg'],
            'radiology': ['x-ray', 'ct scan', 'mri', 'ultrasound', 'radiology'],
            'lab_report': ['laboratory', 'lab results', 'reference range', 'normal']
        }
    
    def setup_ocr(self):
        """Initialize OCR engines"""
        try:
            if ADVANCED_PROCESSING:
                self.ocr_reader = easyocr.Reader(['en'])
                logging.info("EasyOCR initialized successfully")
        except Exception as e:
            logging.warning(f"Advanced OCR not available: {e}")
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Main method to process any supported document type
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text, metadata, and analysis
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        file_size = os.path.getsize(file_path)
        
        result = {
            'file_path': file_path,
            'file_name': Path(file_path).name,
            'file_extension': file_extension,
            'file_size_bytes': file_size,
            'processed_at': datetime.now().isoformat(),
            'text_content': '',
            'metadata': {},
            'document_type': 'unknown',
            'confidence_score': 0.0,
            'error': None
        }
        
        try:
            # Route to appropriate processor
            if file_extension in self.supported_formats['pdf']:
                if PDF_AVAILABLE:
                    result.update(self.process_pdf(file_path))
                else:
                    result['error'] = "PDF processing libraries not available. Please install with: pip install PyPDF2"
            elif file_extension in self.supported_formats['word']:
                if PDF_AVAILABLE:
                    result.update(self.process_word_document(file_path))
                else:
                    result['error'] = "Word processing libraries not available. Please install with: pip install python-docx"
            elif file_extension in self.supported_formats['text']:
                result.update(self.process_text_file(file_path))
            elif file_extension in self.supported_formats['image']:
                if PDF_AVAILABLE:
                    result.update(self.process_image(file_path))
                else:
                    result['error'] = "Image processing libraries not available. Please install with: pip install opencv-python pytesseract"
            else:
                result['error'] = f"Unsupported file format: {file_extension}"
            
            # Classify document type based on content
            if result['text_content'] and not result['error']:
                result['document_type'] = self.classify_document_type(result['text_content'])
                result['confidence_score'] = self.calculate_confidence_score(result)
                
        except Exception as e:
            result['error'] = str(e)
            logging.error(f"Error processing document {file_path}: {e}")
        
        return result
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF documents"""
        result = {'text_content': '', 'metadata': {}}
        
        if not PDF_AVAILABLE:
            raise Exception("PDF processing libraries not available")
        
        try:
            # Try advanced PDF processing first
            if ADVANCED_PROCESSING:
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    
                    result['text_content'] = '\n\n'.join(text_parts)
                    result['metadata'] = {
                        'pages': len(pdf.pages),
                        'method': 'pdfplumber'
                    }
            else:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    
                    result['text_content'] = '\n\n'.join(text_parts)
                    result['metadata'] = {
                        'pages': len(pdf_reader.pages),
                        'method': 'PyPDF2'
                    }
                    
        except Exception as e:
            raise Exception(f"Error processing PDF: {e}")
        
        return result
    
    def process_word_document(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents"""
        result = {'text_content': '', 'metadata': {}}
        
        if not PDF_AVAILABLE:
            raise Exception("Word processing libraries not available")
        
        try:
            if ADVANCED_PROCESSING and file_path.endswith('.docx'):
                # Use mammoth for better formatting preservation
                with open(file_path, "rb") as docx_file:
                    result_mammoth = mammoth.extract_text(docx_file)
                    result['text_content'] = result_mammoth.value
                    result['metadata'] = {
                        'method': 'mammoth',
                        'warnings': len(result_mammoth.messages)
                    }
            else:
                # Fallback to python-docx
                doc = Document(file_path)
                paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                result['text_content'] = '\n\n'.join(paragraphs)
                result['metadata'] = {
                    'paragraphs': len(doc.paragraphs),
                    'method': 'python-docx'
                }
                
        except Exception as e:
            raise Exception(f"Error processing Word document: {e}")
        
        return result
    
    def process_text_file(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        result = {'text_content': '', 'metadata': {}}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                result['text_content'] = content
                result['metadata'] = {
                    'lines': len(content.split('\n')),
                    'characters': len(content),
                    'method': 'direct_read'
                }
                
        except Exception as e:
            raise Exception(f"Error processing text file: {e}")
        
        return result
    
    def process_image(self, file_path: str) -> Dict[str, Any]:
        """Process images using OCR"""
        result = {'text_content': '', 'metadata': {}}
        
        if not PDF_AVAILABLE:
            raise Exception("Image processing libraries not available")
        
        try:
            # Load and preprocess image
            image = cv2.imread(file_path)
            if image is None:
                raise Exception("Could not load image")
            
            # Get image metadata
            height, width = image.shape[:2]
            result['metadata'] = {
                'width': width,
                'height': height,
                'channels': image.shape[2] if len(image.shape) > 2 else 1
            }
            
            # Preprocess image for better OCR
            processed_image = self.preprocess_image_for_ocr(image)
            
            # Try advanced OCR first
            if self.ocr_reader and ADVANCED_PROCESSING:
                try:
                    ocr_results = self.ocr_reader.readtext(processed_image)
                    text_parts = [result[1] for result in ocr_results if result[2] > 0.5]
                    result['text_content'] = ' '.join(text_parts)
                    result['metadata']['method'] = 'easyocr'
                    result['metadata']['confidence'] = sum(r[2] for r in ocr_results) / len(ocr_results) if ocr_results else 0
                except Exception:
                    pass
            
            # Fallback to tesseract
            if not result['text_content']:
                try:
                    # Convert BGR to RGB for PIL
                    rgb_image = cv2.cvtColor(processed_image, cv2.COLOR_BGR2RGB)
                    pil_image = Image.fromarray(rgb_image)
                    
                    # Use tesseract with custom config for medical documents
                    custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,;:!?()[]{}+-=*/% '
                    text = pytesseract.image_to_string(pil_image, config=custom_config)
                    result['text_content'] = text.strip()
                    result['metadata']['method'] = 'tesseract'
                except Exception as e:
                    logging.warning(f"Tesseract OCR failed: {e}")
            
            # If this looks like a medical chart/ECG, try specialized processing
            if self.is_medical_chart(image):
                medical_data = self.extract_medical_chart_data(image)
                result['metadata']['medical_chart_data'] = medical_data
                
        except Exception as e:
            raise Exception(f"Error processing image: {e}")
        
        return result
    
    def preprocess_image_for_ocr(self, image):
        """Preprocess image to improve OCR accuracy"""
        if not PDF_AVAILABLE:
            return image
        
        # Convert to grayscale
        if len(image.shape) > 2:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image
        
        # Apply denoising
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Enhance contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        
        # Apply threshold for better text recognition
        _, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return binary
    
    def is_medical_chart(self, image) -> bool:
        """Detect if image contains medical charts/graphs"""
        if not PDF_AVAILABLE:
            return False
        
        # Simple heuristic: look for grid patterns and regular waves
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) > 2 else image
        
        # Detect lines (common in medical charts)
        edges = cv2.Canny(gray, 50, 150, apertureSize=3)
        lines = cv2.HoughLines(edges, 1, np.pi/180, threshold=100)
        
        return lines is not None and len(lines) > 10
    
    def extract_medical_chart_data(self, image) -> Dict[str, Any]:
        """Extract data from medical charts/ECG"""
        # Placeholder for advanced medical chart analysis
        # This would require specialized medical image processing
        return {
            'type': 'medical_chart',
            'analysis': 'Basic chart detected - advanced analysis requires specialized libraries',
            'recommendation': 'Consider manual review by healthcare professional'
        }
    
    def classify_document_type(self, text: str) -> str:
        """Classify document type based on text content"""
        text_lower = text.lower()
        
        # Count matches for each medical document type
        type_scores = {}
        for doc_type, keywords in self.medical_patterns.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                type_scores[doc_type] = score
        
        if type_scores:
            return max(type_scores, key=type_scores.get)
        
        # Generic classification
        if any(word in text_lower for word in ['patient', 'medical', 'doctor', 'hospital']):
            return 'medical_document'
        elif any(word in text_lower for word in ['report', 'analysis', 'findings']):
            return 'report'
        else:
            return 'general_document'
    
    def calculate_confidence_score(self, result: Dict[str, Any]) -> float:
        """Calculate confidence score for the document processing"""
        score = 0.0
        
        # Base score for successful text extraction
        if result['text_content']:
            score += 0.5
        
        # Bonus for longer text (more content usually means better extraction)
        text_length = len(result['text_content'])
        if text_length > 100:
            score += 0.2
        elif text_length > 50:
            score += 0.1
        
        # Bonus for medical document classification
        if result['document_type'] != 'unknown':
            score += 0.2
        
        # OCR confidence if available
        if 'confidence' in result.get('metadata', {}):
            score += result['metadata']['confidence'] * 0.1
        
        return min(score, 1.0)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of all supported file formats"""
        formats = []
        for format_list in self.supported_formats.values():
            formats.extend(format_list)
        return formats
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate if file can be processed"""
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        file_extension = Path(file_path).suffix.lower()
        if file_extension not in self.get_supported_formats():
            return False, f"Unsupported file format: {file_extension}"
        
        file_size = os.path.getsize(file_path)
        max_size = 50 * 1024 * 1024  # 50MB limit
        if file_size > max_size:
            return False, f"File too large: {file_size / (1024*1024):.1f}MB (max: 50MB)"
        
        return True, "File is valid"